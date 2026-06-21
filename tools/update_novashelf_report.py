# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path.cwd()
SOURCE = ROOT / "2026年计算机综合实训报告-NovaShelf灵感资料工作台.docx"
BACKUP = ROOT / "2026年计算机综合实训报告-NovaShelf灵感资料工作台.修改前备份.docx"
OUTPUT = SOURCE

BODY_HEADER = "华东交通大学计算机综合实训报告（NovaShelf灵感资料工作台）"
TASK_HEADER = "华东交通大学计算机综合实训任务书"
FONT_CN = "宋体"
FONT_EN = "Times New Roman"


REPLACEMENTS = [
    ("Nova shelf", "NovaShelf"),
    ("Nova Shelf", "NovaShelf"),
    ("ACGShare资源分享与评论平台", "NovaShelf灵感资料工作台"),
    ("ACGShare 资源分享与评论平台", "NovaShelf灵感资料工作台"),
    ("ACGShare", "NovaShelf"),
    ("ACG 资源库", "NovaShelf资料库"),
    ("ACG 类学习资料", "创作类学习资料"),
    ("ACG 学习素材", "创作学习素材"),
    ("ACG 资源", "创作资料"),
    ("ACG", "创作素材"),
    ("资源分享与评论平台", "灵感资料工作台"),
    ("资源工作平台", "灵感资料工作台"),
    ("资源库", "资料库"),
    ("资源列表", "资料列表"),
    ("资源详情", "资料详情"),
    ("资源浏览", "资料浏览"),
    ("资源管理", "资料编排"),
    ("资源新增", "资料新增"),
    ("新增资源", "新增资料"),
    ("编辑资源", "编辑资料"),
    ("删除资源", "删除资料"),
    ("资源表单", "资料表单"),
    ("资源标题", "资料标题"),
    ("资源分类", "资料类型"),
    ("资源统计", "资料统计"),
    ("资源筛选", "资料筛选"),
    ("资源下载", "资料入口"),
    ("资源信息", "资料信息"),
    ("资源数据", "资料数据"),
    ("资源维护", "资料维护"),
    ("资源审核", "资料审核"),
    ("资源级联", "资料级联"),
    ("资源总数", "资料总数"),
    ("资源数", "资料数"),
    ("资源数统计", "资料数统计"),
    ("资源表", "资料表"),
    ("资源对象", "资料对象"),
    ("资源", "资料"),
    ("评论管理", "反馈审核"),
    ("评论列表", "反馈列表"),
    ("发表评论", "提交反馈"),
    ("发布评论", "提交反馈"),
    ("删除评论", "删除反馈"),
    ("查看评论", "查看反馈"),
    ("空评论", "空反馈"),
    ("评论内容", "反馈内容"),
    ("评论审核", "反馈审核"),
    ("评论数据", "反馈数据"),
    ("评论数", "反馈数"),
    ("评论", "反馈"),
    ("下载链接", "资料入口"),
    ("下载入口", "资料入口"),
    ("后台资源管理", "后台资料编排"),
    ("后台管理", "素材控制台"),
    ("数据统计", "数据看板"),
    ("Galgame", "互动剧本"),
    ("轻小说", "叙事文本"),
    ("动漫资料", "视觉设定"),
    ("工具资料", "创作工具"),
    ("工具资源", "创作工具"),
    ("音乐资料", "声音素材"),
    ("音乐资源", "声音素材"),
    ("其他", "灵感备忘"),
    ("mode=mysql", "mode=mysql，数据库为 novashelf"),
    ("acgshare", "novashelf"),
    ("acgshare-backend", "nova-shelf-backend"),
    ("acgshare-frontend", "nova-shelf-frontend"),
]


SECTION_TITLE_REPLACEMENTS = {
    "5.1 首页资源浏览与筛选": "5.1 首页资料浏览与筛选",
    "5.2 资源详情、评论与评分": "5.2 资料详情、反馈与评分",
    "5.4 后台资源管理": "5.4 后台资料编排",
    "5.6 资源新增与封面上传": "5.6 资料新增与封面上传",
}


def apply_replacements(text: str) -> str:
    value = text
    for old, new in REPLACEMENTS:
        value = value.replace(old, new)
    # Clean awkward double replacements caused by broad "资源" -> "资料".
    value = value.replace("资料资料", "资料")
    value = value.replace("创作资料学习素材", "创作学习素材")
    value = value.replace("互动剧本、叙事文本、视觉设定、创作工具和声音素材等创作素材", "互动剧本、叙事文本、视觉设定、创作工具和声音素材等资料")
    return value


def replace_paragraph_text(paragraph, text: str):
    if paragraph.text == text:
        return
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return
    for run in runs[1:]:
        paragraph._p.remove(run._r)
    runs[0].text = text


def replace_cell_text(cell, text: str):
    if not cell.paragraphs:
        cell.text = text
        return
    p = cell.paragraphs[0]
    replace_paragraph_text(p, text)
    for extra in list(cell.paragraphs)[1:]:
        parent = extra._element.getparent()
        if parent is not None:
            parent.remove(extra._element)


def set_run_font(run, size=9):
    run.font.name = FONT_CN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run.font.size = Pt(size)


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def add_page_number(paragraph):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = paragraph.add_run("- ")
    set_run_font(left, size=9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), FONT_CN)
    r_fonts.set(qn("w:ascii"), FONT_EN)
    r_fonts.set(qn("w:hAnsi"), FONT_EN)
    r_pr.append(r_fonts)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(r_pr)
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    right = paragraph.add_run(" -")
    set_run_font(right, size=9)


def set_page_start(section, start=None):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if start is None:
        if pg_num_type is not None and pg_num_type.get(qn("w:start")) is not None:
            del pg_num_type.attrib[qn("w:start")]
        return
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def set_header(section, text):
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    clear_paragraph(header)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = header.add_run(text)
        set_run_font(run, size=9)


def clear_footer(section):
    footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    clear_paragraph(footer)


def configure_headers(doc):
    sections = list(doc.sections)
    for idx, section in enumerate(sections, start=1):
        set_page_start(section, None)
        if idx == 1:
            set_header(section, "")
            clear_footer(section)
        elif idx in (2, 3):
            set_header(section, TASK_HEADER)
            clear_footer(section)
        elif idx == 4:
            set_header(section, "")
            clear_footer(section)
        else:
            set_header(section, BODY_HEADER)
            footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
            add_page_number(footer)
            if idx == 5:
                set_page_start(section, 1)


def main():
    if not BACKUP.exists():
        shutil.copyfile(SOURCE, BACKUP)
    doc = Document(SOURCE)

    first_heading_idx = next(
        (i for i, p in enumerate(doc.paragraphs) if p.style.name == "Heading 1" and p.text.strip().startswith("第一章")),
        None,
    )
    # Keep the cover untouched. Start from the task book/catalog/body text.
    edit_start = 9
    changes = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < edit_start:
            continue
        original = paragraph.text
        if not original:
            continue
        updated = SECTION_TITLE_REPLACEMENTS.get(original.strip(), apply_replacements(original))
        if updated != original:
            replace_paragraph_text(paragraph, updated)
            changes += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original = cell.text
                updated = apply_replacements(original)
                if updated != original:
                    replace_cell_text(cell, updated)
                    changes += 1

    configure_headers(doc)
    doc.save(OUTPUT)
    print(json.dumps({"updated": str(OUTPUT), "backup": str(BACKUP), "changes": changes}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
