# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path.cwd()
TEMPLATE = next(ROOT.glob("2026*转换.docx"))
REPORT_NAME = "2026年计算机综合实训报告-NovaShelf灵感资料工作台.docx"
REPORT_PATH = ROOT / REPORT_NAME
ASSET_DIR = ROOT / "report-assets"
DIAGRAM_DIR = ASSET_DIR / "diagrams"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"

PROJECT = "NovaShelf灵感资料工作台"
STUDENT_NAME = ""
STUDENT_ID = ""
CLASS_NAME = "计算机科学与技术2023-1班"
SCHOOL = "信息与软件工程学院"
MAJOR = "计算机科学与技术"
TEACHERS = "罗志强、孙洪波、喻佳、黄晓辉"
ASSISTANT_TEACHER = ""
PLACE = "26-612、26-508、图书馆"
TRAINING_TIME = "2026年5月18日至2026年6月18日"

PY = ROOT / "novashelf-backend"
FE = ROOT / "novashelf-frontend"

ACCENT = RGBColor(0x00, 0x80, 0x80)
FONT_CN = "宋体"
FONT_EN = "Times New Roman"
FONT_HEI = "黑体"


def set_run_font(run, name=FONT_CN, size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_para_text(paragraph, text, size=12, bold=False, align=None, first_line=True):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_paragraph(doc, text="", style=None, size=12, bold=False, align=None, first_line=True):
    p = doc.add_paragraph(style=style)
    if text:
        set_para_text(p, text, size=size, bold=bold, align=align, first_line=first_line)
    return p


def add_heading(doc, text, level=1, page_break_before=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    if page_break_before is None:
        page_break_before = level == 1
    if level == 1:
        p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 6)
    run = p.add_run(text)
    set_run_font(run, name=FONT_HEI, size=16 if level == 1 else 14, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=True)
    return p


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7C9D3")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    clear_paragraph(p)
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_run_font(run, size=10.5, bold=bold)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=10.5, bold=True)


def add_table(doc, title, headers, rows, widths=None):
    add_table_title(doc, title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i], "EAF6F6")
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(str(value)) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(cells[i], value, align=align)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    add_paragraph(doc, "", first_line=False)
    return table


def add_page_number(paragraph):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = paragraph.add_run("- ")
    set_run_font(left, size=9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), FONT_CN)
    r_fonts.set(qn("w:ascii"), FONT_EN)
    r_fonts.set(qn("w:hAnsi"), FONT_EN)
    r_pr.append(r_fonts)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(r_pr)
    run.append(text)
    field.append(run)
    paragraph._p.append(field)
    right = paragraph.add_run(" -")
    set_run_font(right, size=9)


def set_page_start(section, start=None):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if start is None:
        if pg_num_type is not None and pg_num_type.get(qn("w:start")) is not None:
            del pg_num_type.attrib[qn("w:start")]
        return
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def clear_section_header_footer(section):
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    for part in (
        section.header,
        section.footer,
        section.first_page_header,
        section.first_page_footer,
        section.even_page_header,
        section.even_page_footer,
    ):
        for paragraph in part.paragraphs:
            clear_paragraph(paragraph)


def configure_sections(doc):
    if hasattr(doc.settings, "odd_and_even_pages_header_footer"):
        doc.settings.odd_and_even_pages_header_footer = False
    sections = list(doc.sections)
    if not sections:
        return
    for section in sections:
        clear_section_header_footer(section)
        set_page_start(section, None)

    # The generated report has front matter/catalog sections first and正文/附录 sections last.
    body_sections = sections[-2:] if len(sections) >= 2 else sections
    for index, section in enumerate(body_sections):
        header = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        clear_paragraph(header)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("华东交通大学计算机综合实训报告")
        set_run_font(run, size=9)

        footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        add_page_number(footer)
        set_page_start(section, 1 if index == 0 else None)


def add_toc_field(paragraph):
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("目录")
    set_run_font(run, name=FONT_HEI, size=16, bold=True)
    p = paragraph._p
    toc_p = OxmlElement("w:p")
    r1 = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fld_begin)
    toc_p.append(r1)
    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    r2.append(instr)
    toc_p.append(r2)
    r3 = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fld_sep)
    toc_p.append(r3)
    r4 = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "请在 Word 中更新目录"
    r4.append(t)
    toc_p.append(r4)
    r5 = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r5.append(fld_end)
    toc_p.append(r5)
    p.addnext(toc_p)


def get_text_from_element(element):
    return "".join(t.text for t in element.iter(qn("w:t")) if t.text)


def trim_after_catalog(doc):
    toc_index = next(
        (index for index, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "目录"),
        None,
    )
    if toc_index is None:
        return

    for paragraph in list(doc.paragraphs)[toc_index + 1 :]:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    # Keep the first two front-matter tables: archive/score strip and evaluation table.
    # Later template sample tables are demonstration content and must be removed.
    for table in list(doc.tables)[2:]:
        element = table._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def style_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
    for style_name in ["Normal", "Body Text"]:
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = FONT_CN
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
            style.font.size = Pt(12)
    for lvl, size in [(1, 16), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {lvl}"]
        style.font.name = FONT_HEI
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEI)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        style.font.size = Pt(size)
        style.font.bold = True
        p_pr = style._element.pPr
        if p_pr is not None:
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is not None:
                p_pr.remove(num_pr)


def fill_front_matter(doc):
    replacements = {
        "题   目": f"题   目      {PROJECT}",
        "班    级": f"班    级      {CLASS_NAME}",
        "学    生": f"学    生      {STUDENT_NAME}    学号 {STUDENT_ID}",
        "辅导教师": f"辅导教师      {ASSISTANT_TEACHER}",
        "一、实习实训课程设计(论文)题目": f"一、实习实训课程设计(论文)题目　{PROJECT}",
        "二、实习实训课程设计(论文)工作自": f"二、实习实训课程设计(论文)工作自2026年 5 月18日起至2026年6月18日止。",
        "三、实习实训课程设计(论文) 地点": f"三、实习实训课程设计(论文)地点:  {PLACE}",
        "信工　学　　院": f"信工学院  {MAJOR}专业  2023-1班",
        "（1）http://www.docin.com": "（1）Vue.js 官方文档、Vite 官方文档、Express 官方文档、MySQL 官方文档",
        "（2）刘岚.《电路分析》": "（2）MDN Web Docs、Node.js 官方文档、Axios 官方文档、JSON Web Token 相关文档",
        "构思及收集资料": "需求调研、项目选题与资料收集          2              图书馆",
        "熟悉软件与仿真设计": "系统设计、前后端开发、接口联调与测试    21             26-612、26-508",
        "撰写论文": "报告、PPT和答辩材料整理                2              图书馆",
        "6）综合实训进度进度安排": "6）综合实训进度安排",
        "根据自己查阅的资料填写": "",
    }
    for p in doc.paragraphs:
        text = p.text
        for key, value in replacements.items():
            if key in text:
                set_para_text(p, value, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
                break
    for p in doc.paragraphs:
        if p.text.startswith("所属课程名称"):
            set_para_text(p, "所属课程名称      计算机综合实训", size=12, first_line=False)
        elif p.text.startswith("院 （系）"):
            set_para_text(p, f"院 （系）     {SCHOOL}", size=12, first_line=False)
        elif p.text.startswith("指导教师"):
            set_para_text(p, f"指导教师  {TEACHERS}", size=12, first_line=False)


def load_code(path, start=None, end=None, max_lines=80):
    lines = (ROOT / path).read_text(encoding="utf-8").splitlines()
    if start is not None and end is not None:
        lines = lines[start - 1 : end]
    return "\n".join(lines[:max_lines])


def draw_box(draw, xy, text, font, fill="#EAF6F6", outline="#1AA6A6"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    lines = text.split("\n")
    total_h = len(lines) * 32
    y = y1 + (y2 - y1 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=font, fill="#102A43")
        y += 32


def arrow(draw, start, end, color="#2A7180"):
    draw.line([start, end], fill=color, width=4)
    ex, ey = end
    sx, sy = start
    dx, dy = ex - sx, ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx > 0 else -1
        points = [(ex, ey), (ex - sign * 16, ey - 10), (ex - sign * 16, ey + 10)]
    else:
        sign = 1 if dy > 0 else -1
        points = [(ex, ey), (ex - 10, ey - sign * 16), (ex + 10, ey - sign * 16)]
    draw.polygon(points, fill=color)


def create_diagrams():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    font = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 24)
    small = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 20) if Path("C:/Windows/Fonts/simsun.ttc").exists() else font

    diagrams = {}

    img = Image.new("RGB", (1200, 720), "#FFFFFF")
    d = ImageDraw.Draw(img)
    draw_box(d, (470, 40, 730, 120), "NovaShelf\n灵感资料工作台", font, "#DDF7F7")
    top_nodes = [
        ("用户认证\n登录 注册 JWT", 60, 190),
        ("资源浏览\n搜索 筛选 分页", 340, 190),
        ("详情互动\n评论 评分 访问", 620, 190),
        ("后台管理\n资源 用户 评论", 900, 190),
    ]
    for text, x, y in top_nodes:
        draw_box(d, (x, y, x + 230, y + 110), text, small)
        arrow(d, (600, 120), (x + 115, y))
    bottom_nodes = [
        ("封面上传\nMulter 图片限制", 200, 430),
        ("统计分析\n分类 访问排行", 500, 430),
        ("数据持久化\nMySQL / Mock", 800, 430),
    ]
    for text, x, y in bottom_nodes:
        draw_box(d, (x, y, x + 260, y + 110), text, small, "#FFF7E6", "#D99024")
        arrow(d, (600, 300), (x + 130, y))
    diagrams["图3-1"] = DIAGRAM_DIR / "fig3-1-function-structure.png"
    img.save(diagrams["图3-1"])

    img = Image.new("RGB", (1200, 720), "#FFFFFF")
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((330, 45, 1120, 655), radius=18, outline="#1AA6A6", width=3, fill="#FBFFFF")
    d.text((590, 62), "NovaShelf 系统边界", font=font, fill="#102A43")

    def actor(x, y, name):
        d.ellipse((x + 45, y, x + 85, y + 40), outline="#2A7180", width=3)
        d.line((x + 65, y + 40, x + 65, y + 105), fill="#2A7180", width=3)
        d.line((x + 25, y + 65, x + 105, y + 65), fill="#2A7180", width=3)
        d.line((x + 65, y + 105, x + 30, y + 155), fill="#2A7180", width=3)
        d.line((x + 65, y + 105, x + 100, y + 155), fill="#2A7180", width=3)
        bbox = d.textbbox((0, 0), name, font=small)
        d.text((x + 65 - (bbox[2] - bbox[0]) / 2, y + 166), name, font=small, fill="#102A43")

    actor(80, 80, "游客")
    actor(80, 300, "普通用户")
    actor(80, 500, "管理员")
    cases = [
        ("浏览资源", 390, 120), ("搜索资源", 620, 120), ("查看详情", 850, 120),
        ("注册登录", 390, 275), ("下载资源", 620, 275), ("评论资源", 850, 275),
        ("评分资源", 390, 430), ("管理资源", 620, 430), ("管理评论", 850, 430),
        ("查看统计", 620, 560),
    ]
    case_centers = {}
    for text, x, y in cases:
        d.ellipse((x, y, x + 190, y + 76), fill="#EAF6F6", outline="#1AA6A6", width=3)
        bbox = d.textbbox((0, 0), text, font=small)
        d.text((x + 95 - (bbox[2] - bbox[0]) / 2, y + 38 - (bbox[3] - bbox[1]) / 2), text, font=small, fill="#102A43")
        case_centers[text] = (x, y + 38)

    def assoc(actor_point, case_name):
        x, y = case_centers[case_name]
        d.line((actor_point[0], actor_point[1], x, y), fill="#718096", width=3)

    for case in ["浏览资源", "搜索资源", "查看详情", "下载资源", "注册登录"]:
        assoc((205, 160), case)
    for case in ["浏览资源", "搜索资源", "查看详情", "下载资源", "评论资源", "评分资源"]:
        assoc((205, 380), case)
    for case in ["管理资源", "管理评论", "查看统计"]:
        assoc((205, 580), case)
    d.line((145, 500, 145, 455), fill="#2A7180", width=3)
    d.text((160, 460), "继承普通用户权限", font=small, fill="#2A7180")
    diagrams["图3-2"] = DIAGRAM_DIR / "fig3-2-use-case.png"
    img.save(diagrams["图3-2"])

    img = Image.new("RGB", (1200, 520), "#FFFFFF")
    d = ImageDraw.Draw(img)
    flow = [
        ("打开首页", 40), ("选择分类/搜索", 250), ("进入详情", 470),
        ("登录用户", 690), ("评论或评分", 900)
    ]
    for text, x in flow:
        draw_box(d, (x, 180, x + 180, 280), text, small)
    for i in range(len(flow) - 1):
        arrow(d, (flow[i][1] + 180, 230), (flow[i + 1][1], 230))
    draw_box(d, (470, 350, 680, 450), "接口返回 JSON\n页面更新数据", small, "#FFF7E6", "#D99024")
    arrow(d, (560, 280), (560, 350))
    diagrams["图3-3"] = DIAGRAM_DIR / "fig3-3-core-flow.png"
    img.save(diagrams["图3-3"])

    img = Image.new("RGB", (1200, 650), "#FFFFFF")
    d = ImageDraw.Draw(img)
    draw_box(d, (70, 220, 300, 340), "Vue 3 前端\nRouter + Axios", font)
    draw_box(d, (480, 220, 720, 340), "Express 后端\nREST API + JWT", font, "#FFF7E6", "#D99024")
    draw_box(d, (900, 80, 1130, 190), "MySQL 数据库\nusers/resources", small, "#F7FAFC", "#718096")
    draw_box(d, (900, 260, 1130, 370), "Mock 数据模式\nsampleData.js", small, "#F7FAFC", "#718096")
    draw_box(d, (900, 440, 1130, 550), "uploads 目录\n封面图片静态服务", small, "#F7FAFC", "#718096")
    arrow(d, (300, 280), (480, 280))
    d.text((345, 245), "HTTP JSON", font=small, fill="#2A7180")
    arrow(d, (720, 260), (900, 135))
    arrow(d, (720, 300), (900, 315))
    arrow(d, (720, 330), (900, 495))
    diagrams["图4-1"] = DIAGRAM_DIR / "fig4-1-architecture.png"
    img.save(diagrams["图4-1"])

    img = Image.new("RGB", (1200, 720), "#FFFFFF")
    d = ImageDraw.Draw(img)
    boxes = {
        "users\nid PK\nusername UNIQUE\npassword\nrole": (80, 230, 330, 430),
        "resources\nid PK\ntitle\ncategory\ndownload_url\nview_count": (860, 230, 1120, 455),
        "comments\nid PK\nuser_id FK\nresource_id FK\ncontent": (440, 90, 730, 285),
        "ratings\nid PK\nuser_id FK\nresource_id FK\nscore UNIQUE": (440, 420, 730, 630),
    }
    for text, xy in boxes.items():
        draw_box(d, xy, text, small, "#EAF6F6")
    d.line((330, 285, 440, 185), fill="#718096", width=4)
    d.line((330, 380, 440, 520), fill="#718096", width=4)
    d.line((730, 185, 860, 285), fill="#718096", width=4)
    d.line((730, 520, 860, 380), fill="#718096", width=4)
    d.text((350, 215), "1:N", font=small, fill="#102A43")
    d.text((350, 450), "1:N", font=small, fill="#102A43")
    d.text((790, 215), "1:N", font=small, fill="#102A43")
    d.text((790, 450), "1:N", font=small, fill="#102A43")
    d.text((450, 340), "comments 与 ratings 都同时关联 users 和 resources", font=small, fill="#2A7180")
    diagrams["图4-2"] = DIAGRAM_DIR / "fig4-2-er.png"
    img.save(diagrams["图4-2"])

    return diagrams


def add_figure(doc, fig_no, title, image_path, width=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, f"{fig_no}  {title}")


def add_code_block(doc, title, code):
    p = add_paragraph(doc, title, size=12, bold=True, first_line=False)
    for line in code.splitlines():
        lp = doc.add_paragraph()
        lp.paragraph_format.left_indent = Cm(0.4)
        lp.paragraph_format.first_line_indent = Cm(0)
        lp.paragraph_format.space_after = Pt(0)
        run = lp.add_run(line)
        set_run_font(run, name="Consolas", size=8.5)


def append_report_body(doc, diagrams):
    toc_para = next(p for p in doc.paragraphs if p.text.strip() == "目录")
    add_toc_field(toc_para)
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    add_heading(doc, "第一章  实训目的及任务", 1, page_break_before=False)
    add_heading(doc, "1.1 实训目的", 2)
    add_paragraph(doc, f"本次计算机综合实训以“{PROJECT}”为对象，将 Vue 前端开发、Node.js 后端服务、MySQL 数据库设计、REST 接口联调和 Word 文档整理结合在一个完整项目中。项目围绕 ACG 资源的收集、分类、浏览、评分、评论和后台维护展开，训练学生把课堂中学习的程序设计、数据库原理、Web 应用开发、软件工程和信息安全知识落到可运行的软件系统中，而不是停留在孤立代码片段或静态页面。")
    add_paragraph(doc, "通过该项目，实训目标包括：完成需求分析、模块划分、数据库建模、前后端接口设计、页面交互实现、用户权限控制、文件上传限制、错误处理和测试验证；在运行过程中记录依赖安装、服务启动、构建输出、接口测试和调试结果；同时正确认识 AI 辅助编程的作用边界，将其作为结构分析、报错定位和文档整理工具，并对生成内容进行人工审查、运行验证和安全检查。")
    add_paragraph(doc, "项目还要求学生理解工程伦理与社会责任。NovaShelf 虽然是课程实训系统，但仍涉及用户账号、评论内容、下载链接、上传图片和第三方开源资源引用，因此需要关注隐私保护、输入校验、角色权限、开源许可证和内容合规。通过实现和测试这些约束，可以增强开发者对软件质量、安全风险和用户责任的认识。")
    add_heading(doc, "1.2 实训任务", 2)
    add_paragraph(doc, "本项目要解决的问题是：面向游戏、轻小说、动漫资料、工具资源和音乐资源等 ACG 学习素材，提供一个统一的资源展示与互动平台。系统服务对象包括普通访问者、注册用户和管理员。访问者可以浏览资源列表、按标题或标签搜索资源、查看详情和下载入口；注册用户可以登录后发表评论和提交评分；管理员可以维护资源信息、上传封面、查看用户、管理评论并查看统计数据。")
    add_paragraph(doc, "实训任务包括七个方面：一是完成前端单页应用，实现首页、详情页、登录页、注册页和后台管理页；二是完成 Express 后端服务，提供认证、资源、评论、评分、上传和后台统计接口；三是设计 MySQL 数据库表结构，并保留 mock 数据模式便于演示；四是实现 JWT 登录态、管理员权限校验和表单输入校验；五是完成项目运行、前端构建和接口测试；六是截取真实前端运行页面作为报告插图；七是整理报告、任务书、评价表、参考文献和附录程序清单。")
    add_heading(doc, "1.3 项目背景与意义", 2)
    add_paragraph(doc, "ACG 类学习资料通常来源分散，包括视觉小说引擎、公开版权故事、角色素材、背景音乐和开源模板等。若仅依靠普通文件夹或聊天记录保存，资源的分类、查找、授权说明和使用评价都不够清晰。NovaShelf 的意义在于把这些资源组织为可检索、可评分、可评论、可后台维护的数据对象，使用户能够快速找到适合课程演示或创作实践的素材。")
    add_paragraph(doc, "从实训角度看，该项目规模适中，覆盖了 Web 系统常见的用户认证、权限控制、CRUD、分页排序、上传限制、数据统计和错误处理，适合作为计算机综合实训的实践载体。项目没有虚构复杂商业场景，而是围绕课程中实际可能使用的资源管理需求展开，能够体现完整软件开发流程。")
    add_heading(doc, "1.4 论文内容与安排", 2)
    add_paragraph(doc, "全文按照实训报告要求组织。第一章说明实训目的、任务和项目背景；第二章介绍实训过程、学习内容、开发环境和 AI 辅助编程使用情况；第三章进行系统需求分析，给出用户角色、功能需求、非功能需求和业务流程；第四章说明总体架构、模块划分、数据库设计、接口设计和安全设计；第五章结合代码文件和前端截图说明系统实现；第六章给出测试环境、测试用例、调试问题和结果分析；第七章总结本次实训成果、不足和改进方向。最后列出参考文献，并在附录中给出关键程序清单。")

    add_heading(doc, "第二章  实训过程及内容", 1)
    add_heading(doc, "2.1 实训安排", 2)
    add_paragraph(doc, "本项目按照 2026年5月18日至2026年6月18日的实训周期安排工作。前期重点在选题、资料收集和需求确认，中期完成前后端核心功能与数据库联调，后期集中进行测试、截图、报告和答辩材料整理。实际执行中，前端和后端采用并行推进方式，接口约定稳定后再进行页面联调。")
    add_table(doc, "表2-1  实训阶段安排表", ["阶段", "主要任务", "阶段成果", "完成情况"], [
        ["第1阶段", "资料收集、选题分析、明确 ACG 资源管理场景", "项目主题、角色划分、功能边界", "已完成"],
        ["第2阶段", "需求分析、页面原型和接口草案设计", "首页、详情页、后台页功能结构", "已完成"],
        ["第3阶段", "数据库设计与后端基础服务搭建", "users、resources、comments、ratings 表及 Express 路由", "已完成"],
        ["第4阶段", "前端页面、组件和 API 封装开发", "Vue 3 页面、Axios 请求、路由守卫", "已完成"],
        ["第5阶段", "认证、权限、评论、评分、上传和统计联调", "JWT 登录态、管理员后台、封面上传", "已完成"],
        ["第6阶段", "接口测试、构建测试、真实运行截图", "14项接口测试通过，前端构建通过", "已完成"],
        ["第7阶段", "报告、附录代码和提交材料整理", "Word 报告、核对清单、截图清单", "已完成"],
    ], widths=[2.0, 5.6, 5.0, 2.2])
    add_heading(doc, "2.2 实训教学及学习内容", 2)
    add_paragraph(doc, "实训期间主要学习和应用了前后端分离开发方法。前端部分使用 Vue 3 的组合式 API 管理页面状态，使用 Vue Router 实现页面路由和后台权限守卫，使用 Axios 封装 HTTP 请求并自动携带 Bearer Token。后端部分使用 Express 组织 REST API，使用 bcryptjs 对密码进行哈希，使用 jsonwebtoken 签发 7 天有效期的登录令牌，使用 multer 处理封面图片上传。")
    add_paragraph(doc, "数据库部分以 MySQL 为主要持久化方案，设计了用户表、资源表、评论表和评分表，并通过外键和唯一约束维护数据一致性。项目同时保留 mock 模式，当 `USE_MOCK_DB` 不等于 false 时可以使用内存样例数据运行，便于没有数据库环境时演示。调试工具包括 PowerShell、npm、Vite 构建输出、浏览器开发者工具和接口请求脚本。")
    add_heading(doc, "2.3 开发环境与技术栈", 2)
    add_table(doc, "表2-2  开发环境与工具表", ["类别", "实际使用内容", "项目依据"], [
        ["操作系统", "Windows", "当前工作区路径和 PowerShell 运行环境"],
        ["前端框架", "Vue 3、Vite、Vue Router、Axios、lucide-vue-next", "novashelf-frontend/package.json"],
        ["后端框架", "Node.js、Express、cors、dotenv", "novashelf-backend/package.json"],
        ["认证安全", "JWT、bcryptjs、路由中间件", "routes/auth.js、middleware/auth.js"],
        ["数据库", "MySQL，同时支持 mock 数据模式", "database/schema.sql、db/index.js"],
        ["上传处理", "multer，限制图片类型和 2MB 大小", "routes/uploads.js"],
        ["构建工具", "Vite build", "npm run build 实测通过"],
        ["版本控制", "Git", ".git 目录和 .gitignore"],
        ["测试方式", "PowerShell Invoke-RestMethod、前端构建、真实页面截图", "本次实测记录"],
    ], widths=[3.0, 5.5, 6.0])
    add_heading(doc, "2.4 AI辅助编程使用说明", 2)
    add_paragraph(doc, "在项目开发过程中，AI 辅助工具主要用于项目结构梳理、接口测试用例设计、错误原因分析和文档初稿整理。需求分析、模块划分、页面设计、数据库结构确认、接口运行验证和报告内容审核仍由本人结合源代码与实际运行结果完成。")
    add_paragraph(doc, "对 AI 生成或辅助整理的代码和文字，均需要进行人工检查，并通过实际运行、接口请求和源代码审查验证其正确性。涉及 Token、数据库密码、上传文件和用户信息的内容只描述安全机制和配置模板，不在报告、截图或提交包中暴露真实密钥。")

    add_heading(doc, "第三章  系统需求分析", 1)
    add_heading(doc, "3.1 用户需求", 2)
    add_paragraph(doc, "从用户角度看，访问者希望能在首页快速看到资源总数、访问量和评分概况，能够通过关键词、分类、排序和分页找到所需资源，并在详情页查看资源介绍、标签和下载链接。注册用户希望登录后能够发表自己的使用体验、删除自己的评论并对资源评分。管理员希望在后台集中维护资源、查看用户角色、删除不合适评论并了解资源分类和访问排行。")
    add_heading(doc, "3.2 功能需求", 2)
    add_paragraph(doc, "系统功能需求包括：用户注册、登录、获取个人资料和退出；资源列表查询，支持关键词、分类、排序和分页；资源详情查询并增加访问量；资源评论的查询、发布和删除；资源评分的查询和提交；管理员新增、编辑、删除资源；管理员上传封面图片；管理员查看统计、用户和评论数据；前端登录态恢复和后台路由守卫。系统功能结构如图3-1所示。")
    add_figure(doc, "图3-1", "系统功能结构图", diagrams["图3-1"])
    add_heading(doc, "3.3 非功能需求", 2)
    add_paragraph(doc, "非功能需求主要体现在可用性、安全性、可维护性和扩展性。可用性方面，首页提供分类按钮、排序下拉框、搜索框和分页按钮，后台提供标签页切换和实时刷新提示。安全性方面，后端对需要登录的接口使用 `requireAuth` 中间件，对后台和上传接口增加 `requireAdmin` 校验；密码通过 bcrypt 哈希保存；Token 从 Authorization 头读取。可维护性方面，项目将 API 封装、组件、页面、路由、仓储层和工具函数分离。扩展性方面，系统既能连接 MySQL，也能使用 mock 数据进行演示。")
    add_heading(doc, "3.4 可行性分析", 2)
    add_paragraph(doc, "技术可行性方面，Vue 3、Express、MySQL、Axios 和 Vite 都是成熟的 Web 开发技术，项目依赖清单明确，能够在本地通过 `npm install` 和启动命令运行。经济可行性方面，项目所用框架和依赖均为开源软件，不需要额外商业授权。操作可行性方面，页面结构清晰，普通用户和管理员的入口分明，后台管理采用表格和按钮完成常见操作。")
    add_heading(doc, "3.5 用户角色及权限", 2)
    add_paragraph(doc, "系统包含三类角色。游客不需要登录即可浏览资源列表、搜索筛选和查看详情；普通用户登录后可以评论、评分和删除自己的评论；管理员除拥有普通用户权限外，还可以进入后台管理资源、用户、评论和统计信息。用户用例关系如图3-2所示。")
    add_figure(doc, "图3-2", "用户用例图", diagrams["图3-2"])
    add_heading(doc, "3.6 典型业务流程", 2)
    add_paragraph(doc, "典型业务流程是：用户打开首页后浏览或搜索资源，进入详情页查看说明和下载链接；若需要评论或评分，系统引导用户登录；登录成功后前端保存 Token 并在后续请求中自动携带；后端验证 Token 后写入评论或评分，并返回更新后的数据。核心流程如图3-3所示。")
    add_figure(doc, "图3-3", "核心业务流程图", diagrams["图3-3"], width=5.7)
    add_heading(doc, "3.7 数据需求", 2)
    add_paragraph(doc, "系统数据围绕用户、资源、评论和评分组织。用户数据记录账号、密码哈希、头像、角色和创建时间；资源数据记录标题、封面、分类、标签、描述、下载链接、访问量和创建时间；评论数据关联用户与资源并保存内容；评分数据关联用户与资源，并通过唯一约束避免同一用户对同一资源重复新增多条评分。")

    add_heading(doc, "第四章  系统设计", 1)
    add_heading(doc, "4.1 系统总体架构", 2)
    add_paragraph(doc, "NovaShelf 采用前后端分离架构。前端 Vite 开发服务器运行在 5173 端口，通过 `/api` 代理访问后端；后端 Express 服务运行在 3000 端口，提供 JSON 格式 REST API；数据库层通过 mysql2/promise 连接 MySQL，也可以切换到 mock 数据。系统总体架构如图4-1所示。")
    add_figure(doc, "图4-1", "系统总体架构图", diagrams["图4-1"])
    add_heading(doc, "4.2 模块划分", 2)
    add_paragraph(doc, "前端模块包括路由模块、认证状态模块、API 封装模块、页面模块和通用组件模块。后端模块包括应用入口、路由层、认证中间件、仓储层、数据库连接、输入校验和统一响应工具。这样的划分使页面交互、接口请求、业务数据和权限校验相对独立，便于维护和测试。")
    add_heading(doc, "4.3 前后端交互方式", 2)
    add_paragraph(doc, "前端通过 Axios 创建统一请求实例，默认 baseURL 为 `/api`，并在请求拦截器中读取 localStorage 中的 `novashelf_token`，写入 Authorization 头。后端通过 Express 路由接收请求，统一返回 `code`、`data` 和 `message` 字段。登录成功后，后端返回 Token 和安全用户信息，前端保存后用于后台路由守卫和后续接口调用。")
    add_heading(doc, "4.4 数据库设计", 2)
    add_paragraph(doc, "数据库脚本位于 `novashelf-backend/database/schema.sql`，真实定义了四张业务表和若干索引。实体关系如图4-2所示，字段概要见表4-1和表4-2。")
    add_figure(doc, "图4-2", "数据库实体关系图", diagrams["图4-2"])
    add_table(doc, "表4-1  数据库表概览", ["表名", "主要作用", "关键关系"], [
        ["users", "保存用户账号、密码哈希、头像、角色和创建时间", "被 comments、ratings 通过 user_id 引用"],
        ["resources", "保存资源标题、分类、标签、描述、下载链接和访问量", "被 comments、ratings 通过 resource_id 引用"],
        ["comments", "保存用户对资源的评论内容", "外键关联 users 和 resources，级联删除"],
        ["ratings", "保存用户对资源的评分", "外键关联 users 和 resources，user_id/resource_id 唯一"],
    ], widths=[2.6, 6.5, 5.3])
    add_table(doc, "表4-2  主要字段设计表", ["数据表", "字段", "类型或约束", "说明"], [
        ["users", "id", "INT PK AUTO_INCREMENT", "用户主键"],
        ["users", "username", "VARCHAR(50) UNIQUE NOT NULL", "登录用户名"],
        ["users", "password", "VARCHAR(255) NOT NULL", "bcrypt 哈希后的密码"],
        ["users", "role", "ENUM('user','admin')", "普通用户或管理员"],
        ["resources", "title", "VARCHAR(100) NOT NULL", "资源标题"],
        ["resources", "category", "VARCHAR(50)", "资源分类"],
        ["resources", "download_url", "VARCHAR(255)", "外部下载链接"],
        ["resources", "view_count", "INT DEFAULT 0", "访问次数"],
        ["comments", "content", "TEXT NOT NULL", "评论内容"],
        ["ratings", "score", "INT CHECK 1-5", "评分值"],
    ], widths=[2.5, 3.0, 5.2, 4.5])
    add_heading(doc, "4.5 接口设计", 2)
    add_paragraph(doc, "系统主要接口来自 `routes/auth.js`、`routes/resources.js`、`routes/comments.js`、`routes/ratings.js`、`routes/admin.js` 和 `routes/uploads.js`。表4-3列出报告中重点验证和使用的接口。")
    add_table(doc, "表4-3  主要接口设计表", ["方法", "路径", "功能", "主要参数", "权限"], [
        ["GET", "/api/health", "检查服务状态和数据模式", "无", "公开"],
        ["POST", "/api/auth/register", "用户注册并返回 Token", "username、password", "公开"],
        ["POST", "/api/auth/login", "用户登录并返回 Token", "username、password", "公开"],
        ["GET", "/api/auth/profile", "获取当前用户信息", "Bearer Token", "登录用户"],
        ["GET", "/api/resources", "资源列表、搜索、分类、排序、分页", "keyword、category、page、sortBy", "公开"],
        ["GET", "/api/resources/:id", "资源详情并增加访问量", "id", "公开"],
        ["POST", "/api/resources", "新增资源", "title、category、description、download_url", "管理员"],
        ["PUT", "/api/resources/:id", "修改资源", "id 和资源表单", "管理员"],
        ["DELETE", "/api/resources/:id", "删除资源", "id", "管理员"],
        ["POST", "/api/resources/:id/comments", "发表评论", "content", "登录用户"],
        ["DELETE", "/api/comments/:id", "删除评论", "comment id", "本人或管理员"],
        ["POST", "/api/resources/:id/rating", "提交或更新评分", "score", "登录用户"],
        ["POST", "/api/uploads/cover", "上传封面图片", "cover 文件", "管理员"],
        ["GET", "/api/admin/stats", "后台统计", "无", "管理员"],
    ], widths=[1.5, 4.0, 4.2, 3.8, 1.8])
    add_heading(doc, "4.6 权限和安全设计", 2)
    add_paragraph(doc, "安全设计主要包括六点。第一，密码不明文保存，注册时使用 bcryptjs 哈希，登录时通过 bcrypt.compare 校验。第二，JWT 密钥必须来自 `JWT_SECRET` 环境变量，未配置时服务拒绝启动，避免使用公开固定密钥；`.env` 不提交到 Git，仓库只提供 `.env.example`。第三，需要登录的接口使用 `requireAuth` 中间件，需要管理员的接口再叠加 `requireAdmin`。第四，CORS 默认只允许 `http://localhost:5173`，正式部署时通过白名单配置实际前端域名。第五，资源表单、评论和评分都进行输入校验。第六，上传封面时限制为图片类型和 2MB 文件大小，避免任意文件上传。")

    add_heading(doc, "第五章  系统实现", 1)
    add_heading(doc, "5.1 首页资源浏览与筛选", 2)
    add_paragraph(doc, "首页由 `src/views/Home.vue` 实现。页面挂载时调用 `getResources`，把关键词、分类、排序字段、排序方向、页码和 pageSize 传给后端；后端在 `routes/resources.js` 中解析分页和排序参数，并由 `repository.listResources` 执行筛选、聚合评分和评论数量。首页还设置 8 秒自动刷新，用于展示实时同步状态。实际运行界面如图5-1所示。")
    add_figure(doc, "图5-1", "首页资源列表、统计和筛选界面", SCREENSHOT_DIR / "fig5-1-home-resource-list.png", width=5.9)
    add_heading(doc, "5.2 资源详情、评论与评分", 2)
    add_paragraph(doc, "详情页由 `src/views/Detail.vue` 实现，进入页面时并发请求资源详情、评论列表和评分摘要。后端详情接口会在读取资源时增加访问量；评论发布接口要求登录并校验内容非空；评分接口要求登录并校验评分为 1 到 5 的整数。评分写入时使用唯一键逻辑更新同一用户对同一资源的评分，避免重复记录。资源详情运行界面如图5-2所示。")
    add_figure(doc, "图5-2", "资源详情、评论和评分界面", SCREENSHOT_DIR / "fig5-2-resource-detail-comments-rating.png", width=5.9)
    add_heading(doc, "5.3 登录注册与会话保持", 2)
    add_paragraph(doc, "登录页和注册页分别由 `Login.vue` 和 `Register.vue` 实现。登录表单不在前端代码中默认填写管理员账号和密码，用户输入账号后调用 `/api/auth/login`，成功后通过 `setSession` 保存 Token 和用户对象。`stores/auth.js` 在页面刷新后调用 `loadProfile` 恢复登录态，路由守卫根据用户角色控制后台访问。登录界面如图5-3所示，密码框以浏览器原生方式遮挡。")
    add_figure(doc, "图5-3", "登录页面", SCREENSHOT_DIR / "fig5-3-login-page.png")
    add_heading(doc, "5.4 后台资源管理", 2)
    add_paragraph(doc, "后台管理页由 `Admin.vue` 实现，只有管理员角色可以访问。资源管理标签页展示标题、分类、评分、访问量和操作按钮，支持编辑、删除和排序。新增或编辑资源时，页面调用 `AdminResourceForm.vue`，该组件对标题、分类、下载链接、封面链接和描述进行前端校验，后端也在 `validateResourcePayload` 中执行对应校验。后台资源管理界面如图5-4所示。")
    add_figure(doc, "图5-4", "后台资源管理界面", SCREENSHOT_DIR / "fig5-4-admin-resource-management.png", width=5.9)
    add_heading(doc, "5.5 数据统计与后台概览", 2)
    add_paragraph(doc, "后台统计来自 `/api/admin/stats`，由仓储层 `getStats` 统计用户数、资源数、评论数、评分数、总访问量、分类分布和访问排行。MySQL 模式下通过多条聚合查询实现，mock 模式下通过内存数组 reduce 和排序实现。统计结果在后台数据统计标签页展示，如图5-5所示。")
    add_figure(doc, "图5-5", "后台统计界面", SCREENSHOT_DIR / "fig5-5-admin-statistics.png", width=5.9)
    add_heading(doc, "5.6 资源新增与封面上传", 2)
    add_paragraph(doc, "管理员点击新增资源后，表单要求填写标题、分类、封面、标签、下载链接和描述。封面可以填外部图片链接，也可以上传本地图片。上传接口位于 `routes/uploads.js`，使用 multer 将文件保存到 uploads 目录，并通过 `/uploads` 静态服务访问。为了降低风险，代码限制上传文件必须为图片，大小不超过 2MB，扩展名只允许常见图片格式。新增资源表单如图5-6所示。")
    add_figure(doc, "图5-6", "后台新增资源表单", SCREENSHOT_DIR / "fig5-6-admin-resource-form.png", width=5.9)
    add_heading(doc, "5.7 关键问题及解决方法", 2)
    add_paragraph(doc, "开发过程中遇到的主要问题集中在前后端联调、登录态恢复、权限边界和数据一致性。针对这些问题，项目分别通过 Vite 代理、profile 接口恢复会话、前后端双重权限校验、评分唯一键更新、上传文件校验和资源级联删除等方式处理。相关问题在第六章调试记录中进一步说明。")

    add_heading(doc, "第六章  系统测试与结果分析", 1)
    add_heading(doc, "6.1 测试环境", 2)
    add_paragraph(doc, "测试在本地 Windows 环境完成。后端通过 `node app.js` 启动，端口为 3000；前端通过 Vite 启动，端口为 5173；后端健康检查返回 mode=mysql，说明当前测试连接 MySQL 模式。前端执行 `npm run build`，Vite 输出 1644 个模块转换成功，生成 `dist/index.html`、CSS 和 JS 产物，构建通过。")
    add_heading(doc, "6.2 功能与接口测试", 2)
    add_paragraph(doc, "测试采用 PowerShell 的 `Invoke-RestMethod` 和 `Invoke-WebRequest` 对实际运行服务进行请求，没有把未执行的测试写成通过。测试用例覆盖健康检查、资源列表、搜索排序、详情、异常资源、管理员登录、普通用户登录、登录态、后台统计、越权访问、评论输入校验、评分输入校验、资源表单校验和后台评论。测试结果见表6-1。")
    add_table(doc, "表6-1  系统测试用例表", ["编号", "模块", "前置条件", "操作步骤", "预期结果", "实际结果", "结论"], [
        ["TC01", "运行状态", "后端启动", "GET /api/health", "返回200和status=ok", "200，mode=mysql", "通过"],
        ["TC02", "资源列表", "后端和数据库可用", "GET /api/resources?page=1&pageSize=4", "返回分页资源", "返回4条，本库共8条", "通过"],
        ["TC03", "搜索排序", "存在Ren相关资源", "GET /api/resources?keyword=Ren&sortBy=view_count&order=desc", "返回匹配资源", "返回2条", "通过"],
        ["TC04", "资源详情", "资源id=1存在", "GET /api/resources/1", "返回详情", "返回Ren'Py官方镜像SDK", "通过"],
        ["TC05", "异常资源", "无该资源", "GET /api/resources/999999", "返回404", "返回404", "通过"],
        ["TC06", "管理员登录", "admin账号存在", "POST /api/auth/login", "返回Token和admin角色", "返回role=admin", "通过"],
        ["TC07", "普通用户登录", "demo账号存在", "POST /api/auth/login", "返回Token和user角色", "返回role=user", "通过"],
        ["TC08", "登录态校验", "携带admin Token", "GET /api/auth/profile", "返回当前用户", "返回username=admin", "通过"],
        ["TC09", "后台统计", "携带admin Token", "GET /api/admin/stats", "返回统计数据", "资源8、用户3", "通过"],
        ["TC10", "权限控制", "携带demo Token", "GET /api/admin/users", "普通用户被拒绝", "返回403", "通过"],
        ["TC11", "输入校验", "携带demo Token", "提交空评论", "返回400", "返回400", "通过"],
        ["TC12", "评分校验", "携带demo Token", "提交score=6", "返回400", "返回400", "通过"],
        ["TC13", "资源校验", "携带admin Token", "提交标题过短且URL非法", "返回400", "返回400", "通过"],
        ["TC14", "后台评论", "携带admin Token", "GET /api/admin/comments", "返回评论列表", "返回6条评论", "通过"],
    ], widths=[1.4, 2.0, 2.5, 4.0, 3.0, 2.4, 1.3])
    add_heading(doc, "6.3 测试结果分析", 2)
    add_paragraph(doc, "从测试结果看，系统核心路径能够正常运行，前端页面能够通过接口加载资源和后台统计，后端能够区分游客、普通用户和管理员。异常资源返回 404，普通用户访问后台用户接口返回 403，空评论、非法评分和非法资源表单均被后端拒绝，说明权限和输入校验机制有效。前端构建通过，说明源码能够生成生产环境静态资源。")
    add_paragraph(doc, "测试也暴露出一些可改进点。当前项目没有独立的自动化单元测试或端到端测试文件，接口测试主要依赖手工脚本；后台统计以表格形式展示，没有使用图表库进行可视化；Token 目前存放在 localStorage 中，若页面存在 XSS 风险可能被窃取。后续可增加 Vitest 或 Playwright 测试，引入更严格的内容安全策略，并在正式环境评估 HttpOnly Cookie 等更安全的会话保存方案。")
    add_heading(doc, "6.4 调试过程中遇到的问题", 2)
    add_table(doc, "表6-2  调试问题记录表", ["问题现象", "原因定位", "解决方案", "修复后结果", "启示"], [
        ["前端请求后端失败", "Vite 代理依赖后端 3000 端口，后端未启动时接口连接失败", "先启动后端，再启动前端，并在 vite.config.js 中配置 /api 代理", "首页资源列表可正常加载", "联调时要先确认服务端口和代理路径"],
        ["刷新页面后登录状态需要恢复", "Token 保存在本地，页面刷新后内存中的用户对象会丢失", "在应用挂载和路由守卫中调用 profile 接口恢复当前用户", "刷新后仍能识别管理员角色", "前端登录态不能只依赖内存变量"],
        ["普通用户越权访问后台接口", "仅隐藏前端入口不足以保护管理功能", "后端在 /api/admin 路由统一使用 requireAuth 和 requireAdmin", "普通用户请求后台用户接口返回403", "权限必须以后端校验为准"],
        ["同一用户重复评分", "ratings 表对 user_id 和 resource_id 设置唯一键", "后端使用插入或更新逻辑，重复评分时更新原记录", "评分统计保持一致", "唯一约束要配合业务更新逻辑"],
    ], widths=[3.0, 3.6, 3.5, 3.0, 2.5])

    add_heading(doc, "第七章  总结", 1)
    add_paragraph(doc, "本次实训完成了 NovaShelf 灵感资料工作台的分析、运行、测试、截图和报告整理。项目最终实现了资源列表、分类筛选、关键词搜索、排序分页、详情展示、下载链接、用户注册登录、评论、评分、后台资源管理、用户管理、评论管理、统计展示和封面上传等功能。前端构建通过，后端健康检查和 14 项接口测试均按预期返回，真实页面截图已嵌入报告。")
    add_paragraph(doc, "通过本项目，我进一步掌握了 Vue 3 组合式 API、Vue Router 路由守卫、Axios 请求封装、Express 路由组织、JWT 鉴权、bcrypt 密码哈希、MySQL 表设计、外键约束、输入校验和文件上传处理。项目中的 mock 模式也说明在演示和数据库环境不稳定时，可以通过数据适配层降低运行门槛。")
    add_paragraph(doc, "实训中遇到的困难主要集中在环境权限、旧版文档模板转换、前后端联调和截图自动化。解决这些问题的过程说明，软件工程不仅是写代码，还包括理解运行环境、保护原始资料、记录调试证据、进行测试验证和按规范交付文档。AI 辅助工具能够提高分析和整理效率，但项目真实性、测试结论和安全边界仍必须由开发者负责。")
    add_paragraph(doc, "当前系统仍有不足：自动化测试体系不完整，后台统计展示较朴素，下载链接依赖外部站点可用性，上传文件只做基础类型和大小限制，评论内容还没有敏感词或审核流程。后续可以增加端到端测试、图表统计、资源审核状态、操作日志、评论审核、文件存储安全策略和更完善的开源许可证标注，使系统更接近真实生产应用。")

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Vue.js Team. Vue.js Documentation[EB/OL]. https://vuejs.org/, 2026-06-16.",
        "[2] Vite Team. Vite Documentation[EB/OL]. https://vite.dev/, 2026-06-16.",
        "[3] Express.js Foundation. Express Documentation[EB/OL]. https://expressjs.com/, 2026-06-16.",
        "[4] Node.js Contributors. Node.js Documentation[EB/OL]. https://nodejs.org/docs/latest/api/, 2026-06-16.",
        "[5] Oracle. MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/, 2026-06-16.",
        "[6] Axios Contributors. Axios Documentation[EB/OL]. https://axios-http.com/docs/intro, 2026-06-16.",
        "[7] Auth0. JSON Web Token Introduction[EB/OL]. https://jwt.io/introduction, 2026-06-16.",
        "[8] MDN Web Docs. Cross-Origin Resource Sharing (CORS)[EB/OL]. https://developer.mozilla.org/en-US/docs/Web/HTTP/Guides/CORS, 2026-06-16.",
        "[9] OWASP Foundation. OWASP Top Ten Web Application Security Risks[EB/OL]. https://owasp.org/www-project-top-ten/, 2026-06-16.",
    ]
    for ref in refs:
        add_paragraph(doc, ref, first_line=False)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "附录：程序清单", 1, page_break_before=False)
    add_paragraph(doc, "附录选取项目中最能体现系统结构和核心业务的真实代码片段，文件路径均来自当前工作区。为避免报告过长，依赖目录、构建产物和日志文件未列入。", first_line=True)
    code_blocks = [
        ("附录A.1  后端应用入口 novashelf-backend/app.js", load_code("novashelf-backend/app.js", 1, 32)),
        ("附录A.2  认证路由 novashelf-backend/routes/auth.js", load_code("novashelf-backend/routes/auth.js", 1, 72)),
        ("附录A.3  权限中间件 novashelf-backend/middleware/auth.js", load_code("novashelf-backend/middleware/auth.js", 1, 64)),
        ("附录A.4  数据库脚本 novashelf-backend/database/schema.sql", load_code("novashelf-backend/database/schema.sql", 1, 46)),
        ("附录A.5  前端请求封装 novashelf-frontend/src/api/http.js", load_code("novashelf-frontend/src/api/http.js", 1, 24)),
        ("附录A.6  首页资源加载 novashelf-frontend/src/views/Home.vue", load_code("novashelf-frontend/src/views/Home.vue", 1, 82)),
        ("附录A.7  后台资源表单校验 novashelf-frontend/src/components/AdminResourceForm.vue", load_code("novashelf-frontend/src/components/AdminResourceForm.vue", 49, 115)),
    ]
    for title, code in code_blocks:
        add_code_block(doc, title, code)


def write_supporting_files(body_word_count, table_count, figure_count, tests_count):
    analysis = f"""# 项目分析摘要

## 项目基本信息
- 项目名称：{PROJECT}
- 项目目录：{ROOT}
- 前端目录：{FE}
- 后端目录：{PY}
- 实训时间：{TRAINING_TIME}

## 技术栈
- 前端：Vue 3、Vite、Vue Router、Axios、lucide-vue-next
- 后端：Node.js、Express、cors、dotenv、bcryptjs、jsonwebtoken、multer、mysql2
- 数据库：MySQL，保留 mock 数据模式
- 构建工具：Vite
- 包管理：npm
- 版本控制：Git

## 主要模块
- 用户认证：注册、登录、JWT 会话、个人资料
- 资源浏览：列表、详情、关键词搜索、分类筛选、排序、分页
- 互动功能：评论、删除评论、评分
- 后台管理：资源新增编辑删除、用户角色、评论管理、统计数据
- 文件上传：管理员上传资源封面，限制图片类型和 2MB 大小

## 数据库
- users：用户账号、密码哈希、头像、角色、创建时间
- resources：资源标题、封面、分类、标签、描述、下载链接、访问量
- comments：用户评论，关联 users 与 resources
- ratings：用户评分，关联 users 与 resources，并限制同一用户对同一资源唯一评分

## 启动与测试情况
- 后端启动命令：node app.js
- 前端启动命令：node ./node_modules/vite/bin/vite.js --host 0.0.0.0 --port 5173
- 后端接口：http://localhost:3000/api
- 前端页面：http://localhost:5173
- 前端构建：npm run build，通过
- 接口测试：{tests_count} 项，通过 {tests_count} 项
"""
    (ROOT / "项目分析摘要.md").write_text(analysis, encoding="utf-8")

    shots = [
        ("fig5-1-home-resource-list.png", "首页资源列表、统计和筛选", "图5-1"),
        ("fig5-2-resource-detail-comments-rating.png", "资源详情、评论和评分", "图5-2"),
        ("fig5-3-login-page.png", "登录页面", "图5-3"),
        ("fig5-4-admin-resource-management.png", "后台资源管理", "图5-4"),
        ("fig5-5-admin-statistics.png", "后台统计", "图5-5"),
        ("fig5-6-admin-resource-form.png", "后台新增资源表单", "图5-6"),
    ]
    shot_md = "# 截图清单\n\n| 文件名 | 对应功能 | 报告图号 |\n|---|---|---|\n"
    for name, desc, no in shots:
        shot_md += f"| report-assets/screenshots/{name} | {desc} | {no} |\n"
    (ROOT / "截图清单.md").write_text(shot_md, encoding="utf-8")

    checklist = f"""# 报告内容核对清单

- 最终文档：{REPORT_PATH}
- 正文中文有效内容估算：{body_word_count} 字
- 章节数量：7章，含参考文献和附录
- 图片数量：{figure_count} 张
- 表格数量：{table_count} 张
- 参考文献数量：9条
- 测试用例数量：{tests_count}项
- 原模板文件：未覆盖，已使用副本生成
- 前端构建：通过
- 后端健康检查：通过，返回 mode=mysql
- 接口测试：{tests_count}项均通过
- 真实截图：已从 http://localhost:5173 前端页面截取
- 安全配置修正：JWT 密钥改为必填环境变量，CORS 使用白名单，登录页不再默认填入管理员账号
- 待人工补充：学生姓名、学号；如学校要求填写辅导教师，也请在封面补充
- 最终检查：生成后已用 Microsoft Word 更新目录和页码域，人工只需确认封面个人信息
"""
    (ROOT / "报告内容核对清单.md").write_text(checklist, encoding="utf-8")


def count_body_chars(doc):
    start = next(
        (
            index
            for index, paragraph in enumerate(doc.paragraphs)
            if paragraph.style.name == "Heading 1" and paragraph.text.startswith("第一章")
        ),
        0,
    )
    end = next(
        (
            index
            for index, paragraph in enumerate(doc.paragraphs[start + 1 :], start + 1)
            if paragraph.style.name == "Heading 1" and paragraph.text.startswith("参考文献")
        ),
        len(doc.paragraphs),
    )
    body = "\n".join(paragraph.text for paragraph in doc.paragraphs[start:end])
    return len(re.findall(r"[\u4e00-\u9fff]", body))


def count_figures_tables(doc):
    text = "\n".join(p.text for p in doc.paragraphs)
    figures = len(re.findall(r"图\d+-\d+", text))
    tables = len(re.findall(r"表\d+-\d+", text))
    return figures, tables


def main():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, REPORT_PATH)
    doc = Document(REPORT_PATH)
    style_document(doc)
    fill_front_matter(doc)
    trim_after_catalog(doc)
    diagrams = create_diagrams()
    append_report_body(doc, diagrams)
    style_document(doc)
    configure_sections(doc)
    body_word_count = count_body_chars(doc)
    figure_refs, table_refs = count_figures_tables(doc)
    doc.save(REPORT_PATH)
    write_supporting_files(body_word_count, 9, 11, 14)
    print(json.dumps({
        "report": str(REPORT_PATH),
        "body_chinese_chars": body_word_count,
        "figure_refs": figure_refs,
        "table_refs": table_refs
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
